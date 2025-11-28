import re
from datetime import datetime
from docx import Document
import os

def extract_date_from_text(text: str):
    """
    Extracts the first valid date found in the text.
    Supports formats:
    - YYYY年MM月DD日
    - YYYY/MM/DD
    - YYYY-MM-DD
    - YYYY.MM.DD
    """
    # Regex patterns for common date formats
    patterns = [
        r'(\d{4})年(\d{1,2})月(\d{1,2})日',  # 2023年5月12日
        r'(\d{4})/(\d{1,2})/(\d{1,2})',      # 2023/05/12
        r'(\d{4})-(\d{1,2})-(\d{1,2})',      # 2023-05-12
        r'(\d{4})\.(\d{1,2})\.(\d{1,2})'     # 2023.05.12
    ]

    for pattern in patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            try:
                year, month, day = map(int, match)
                # Basic validation
                if 1900 <= year <= 2100 and 1 <= month <= 12 and 1 <= day <= 31:
                    return datetime(year, month, day)
            except ValueError:
                continue
    
    return None

def extract_date_from_docx(file_path: str):
    """
    Reads a DOCX file and extracts the first date found.
    """
    try:
        doc = Document(file_path)
        full_text = []
        
        # Check paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Check tables (sometimes dates are in headers/tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        text_content = "\n".join(full_text)
        return extract_date_from_text(text_content)
    except Exception as e:
        print(f"Error extracting date from {file_path}: {e}")
        return None
